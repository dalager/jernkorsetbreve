
import json
from MLStripper import strip_tags
from datetime import datetime
from docx import Document
import csv
from ebooklib import epub

import locale
locale.setlocale(locale.LC_TIME, "da_DK") 

def load():
    with open('content.json','r') as json_file:
        letters = json.load(json_file)

        # stripping html
        for l in letters:
            l['Text'] = strip_tags(l['Text'])[2::]
            l['Text'] = l['Text'].replace('\r\n','\n')
            l['Text'] = l['Text'].replace('\n\n','!x')
            l['Text'] = l['Text'].replace('\n',' ')
            l['Text'] = l['Text'].replace('!x','\n')
            l['Text'] = l['Text'].replace('\xad','')
            l['LetterDate'] = datetime.strptime(l['LetterDate'],"%Y-%m-%dT%H:%M:%S")
            l['LetterHeading'] = l['LetterDate'].strftime("%A d. %d. %b %Y").capitalize() + ' - ' + l['Place']
            if(l['Location']):
                l['Location'] = ','.join(l['Location'].split(',')[0:2])
        return letters


def create_document(letters):
    document = Document()
    document.core_properties.author = "Christian Dalager, christian@dalager.com"
    document.core_properties.comments = "Genereret udfra datafil, kan fremsendes"
    document.add_heading('Jernkorset, Breve fra 1. verdenskrig', 0)

    for letter in letters:
        document.add_heading(letter['LetterDate'].strftime("%A d. %d. %b %Y").capitalize() + ' - ' + letter['Place'], level=1)
        p = document.add_paragraph()
        p.add_run('Fra: ').bold = True
        p.add_run(letter['Sender'])
        p.add_run('\n')
        p.add_run('Til: ').bold = True
        p.add_run(letter['Recipient'])
        paragraphs = letter['Text'].split('\n')
        for pg in paragraphs:
            document.add_paragraph(pg)


    document.save('Jernkorset-en-brevsamling.docx')

def save_csv(letters):
    with open('names.csv', 'w', newline='') as csvfile:
        fieldnames = ['placename', 'location','date','sender','recipient','wordcount']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for letter in letters:
            writer.writerow({
                'placename': letter['Place'],
                'location': letter['Location'],
                'date': letter['LetterDate'],
                'sender': letter['Sender'],
                'recipient': letter['Recipient'],
                'wordcount':len(letter['Text'].split())
                })


def create_epub(letters):
    book = epub.EpubBook()
    # set metadata
    book.set_identifier('Jernkorset2020-06-09')
    book.set_title('Jernkorset')
    book.set_language('da')
    book.add_author('Jørgen Dalager')
    book.add_author('Christian Dalager')
    chaps = []

    c1 = epub.EpubHtml(title='Introduktion', file_name='intro.xhtml', lang='da')
    c1.content=u'<html><head></head><body><h1>Jernkorset</h1><p>Denne brevsamling består af 666 breve fra perioden 1911 til 1918, primært fra men også til Peter Mærsk, der under første verdenskrig kæmpede på tysk side som en del af det danske mindretal i sønderjylland.</p></body></html>'
    book.add_item(c1)

    
    for i,letter in enumerate(letters):
        c = epub.EpubHtml(title = letter['LetterHeading'],file_name=f'chap_{i+1}.xhtml')
        html = f'<h1>{letter["LetterHeading"]}</h1>'
        html = html + f'<p>Fra: {letter["Sender"]}<br/>Til: {letter["Recipient"]}</p>'
        html = html + ''.join([f'<p>{p}</p>' for p in letter['Text'].split('\n')])
        if(letter['Location']):
            html = html + f'<p><a href="https://www.google.com/maps/@{letter["Location"]},11z">Se {letter["Place"]} på Google Maps</a></p>'
        c.content = html
        book.add_item(c)
        chaps.append(c)

    book.toc = (epub.Link('intro.xhtml', 'Introduktion', 'intro'),
                 (epub.Section('Brevene'),
                 (tuple(chaps)))
                )


    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    style = 'BODY {color: white;}'
    nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)

    # add CSS file
    book.add_item(nav_css)

    # basic spine
    book.spine = ['nav',c1] + chaps

    epub.write_epub('Jernkorset.epub', book, {})

letters = load()
letters = sorted(letters, key = lambda l:l['LetterDate'])

#create_document(letters)
#save_csv(letters)
create_epub(letters)
